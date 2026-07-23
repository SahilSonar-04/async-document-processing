import time
import json
import uuid
import os
import re
import logging
import chardet
from collections import Counter
from datetime import datetime, timezone
from celery import Task
from celery.exceptions import MaxRetriesExceededError
from langdetect import detect, DetectorFactory, LangDetectException
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session

from app.workers.celery_app import celery_app
from app.core.config import settings
from app.db.redis_client import publish_event_sync
from app.models.models import Job, ProcessingResult, JobStatus, Document

logger = logging.getLogger(__name__)

DetectorFactory.seed = 0

sync_engine = create_engine(
    settings.sync_database_url,
    pool_pre_ping=True,
    pool_size=5,
    pool_recycle=300,
)
SyncSession = sessionmaker(bind=sync_engine, autoflush=False, autocommit=False)

_STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with",
    "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does", "did",
    "will", "would", "could", "should", "may", "might", "must", "can", "shall",
    "this", "that", "these", "those", "it", "its", "i", "you", "he", "she", "we", "they",
    "my", "your", "our", "their", "his", "her", "them", "him", "us", "me",
    "as", "by", "from", "into", "about", "than", "then", "so", "if", "not", "no",
    "up", "down", "out", "off", "over", "under", "again", "further", "here", "there",
    "when", "where", "why", "how", "all", "any", "both", "each", "few", "more",
    "most", "other", "some", "such", "only", "own", "same", "too", "very", "just",
    "also", "one", "two", "first", "second",
}

_BULLET_LINE_RE = re.compile(r"^\s*(?:[\u2022\u25CF\u25AA\u2023\u25E6\u00B7]|[-\u2013\u2014](?=\s))\s*")
_MAX_PHRASE_WORDS = 5
_MIN_STANDALONE_SENTENCE_WORDS = 3
_FRAGMENT_WORD_LIMIT = 4


def _emit(job_id: str, event: str, progress: int, stage: str, message: str = "") -> None:
    payload = {
        "job_id": job_id,
        "event": event,
        "progress": progress,
        "stage": stage,
        "message": message,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }
    try:
        publish_event_sync(job_id, payload)
    except Exception as e:
        logger.warning("_emit failed for job %s: %s", job_id, e)


def _update_job(session: Session, job_id: str, **kwargs) -> Job:
    job = session.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise ValueError(f"Job {job_id} not found in database")
    for k, v in kwargs.items():
        setattr(job, k, v)
    session.commit()
    return job


def _restore_file_from_db_backup(session: Session, path: str, document_id: str) -> bool:
    document = session.query(Document).filter(Document.id == document_id).first()
    if not document or not document.file_content:
        return False

    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "wb") as f:
            f.write(document.file_content)
        logger.info("Restored file from DB backup for document %s -> %s", document_id, path)
        return True
    except Exception as e:
        logger.error("Failed to restore file from DB backup for document %s: %s", document_id, e)
        return False


def _extract_text_from_file(
    path: str, file_type: str, session: Session, document_id: str
) -> str:
    if not os.path.exists(path):
        restored = _restore_file_from_db_backup(session, path, document_id)
        if not restored:
            logger.warning(
                "File not found at %s and no DB backup available — using placeholder text",
                path,
            )
            filename = os.path.basename(path)
            return (
                f"[File unavailable: {filename}]\n\n"
                "The original file could not be read because the server restarted "
                "after upload (ephemeral disk) and no backup was found. For reliable "
                "file processing in production, configure an external object store "
                "(S3, Cloudinary, etc.)."
            )

    try:
        if file_type in ("txt", "md", "csv", "json"):
            with open(path, "rb") as f:
                raw = f.read()
            encoding = chardet.detect(raw)["encoding"] or "utf-8"
            return raw.decode(encoding, errors="replace")

        elif file_type == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return raw_text.strip()

        elif file_type == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts)

        else:
            filename = os.path.basename(path)
            return f"[Unsupported file type: {file_type} for {filename}]"

    except Exception as e:
        return f"[Error extracting text: {str(e)}]"


def _normalize_text(text: str) -> str:
    units: list[str] = []
    buffer = ""

    def flush():
        nonlocal buffer
        if buffer:
            u = buffer.strip()
            if u and u[-1] not in ".!?:;":
                u += "."
            units.append(u)
        buffer = ""

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        without_bullet = _BULLET_LINE_RE.sub("", line)
        is_bullet_start = without_bullet != line
        content = without_bullet.strip()
        word_count = len(content.split())

        if is_bullet_start:
            flush()
            buffer = content
            continue

        if buffer and word_count <= _FRAGMENT_WORD_LIMIT and buffer[-1:] not in ".!?:;":
            buffer = f"{buffer} {content}".strip()
        elif buffer:
            flush()
            buffer = content
        else:
            buffer = content

    flush()
    return " ".join(units)


def _split_sentences(text: str) -> list[str]:
    raw_parts = [p.strip() for p in re.split(r"(?<=[.!?:])\s+", text.strip()) if p.strip()]

    merged: list[str] = []
    for part in raw_parts:
        if merged and len(part.split()) < _MIN_STANDALONE_SENTENCE_WORDS:
            merged[-1] = f"{merged[-1]} {part}"
        else:
            merged.append(part)
    return merged


def _summarize(text: str, max_sentences: int = 3) -> str:
    sentences = _split_sentences(text)
    if len(sentences) <= max_sentences:
        return " ".join(sentences)

    words = re.findall(r"[a-zA-Z']+", text.lower())
    freq = Counter(w for w in words if w not in _STOPWORDS and len(w) > 2)
    if not freq:
        return " ".join(sentences[:max_sentences])

    max_freq = max(freq.values())
    for w in freq:
        freq[w] /= max_freq

    scored = []
    for i, sentence in enumerate(sentences):
        sent_words = re.findall(r"[a-zA-Z']+", sentence.lower())
        if len(sent_words) < _MIN_STANDALONE_SENTENCE_WORDS:
            continue
        score = sum(freq.get(w, 0) for w in sent_words) / len(sent_words)
        score *= 1.0 - 0.05 * min(i, 5)
        scored.append((score, i, sentence))

    if not scored:
        return " ".join(sentences[:max_sentences])

    top = sorted(scored, key=lambda x: x[0], reverse=True)[:max_sentences]
    ordered = [s for _, _, s in sorted(top, key=lambda x: x[1])]
    return " ".join(ordered)


def _rake_keywords(text: str, max_keywords: int = 10) -> list[str]:
    sentences = _split_sentences(text)
    if not sentences:
        return []

    phrases: list[list[str]] = []
    for sentence in sentences:
        words = re.findall(r"[a-zA-Z][a-zA-Z'-]*", sentence.lower())
        current: list[str] = []
        for w in words:
            if w in _STOPWORDS or len(w) <= 2:
                if current:
                    phrases.append(current)
                    current = []
            else:
                current.append(w)
                if len(current) >= _MAX_PHRASE_WORDS:
                    phrases.append(current)
                    current = []
        if current:
            phrases.append(current)

    if not phrases:
        return []

    freq: Counter = Counter()
    degree: Counter = Counter()
    for phrase in phrases:
        co_degree = len(phrase) - 1
        for w in phrase:
            freq[w] += 1
            degree[w] += co_degree
    for w in freq:
        degree[w] += freq[w]

    word_score = {w: degree[w] / freq[w] for w in freq}

    seen: set[str] = set()
    phrase_scores: list[tuple[float, str]] = []
    for phrase in phrases:
        key = " ".join(phrase)
        if key in seen:
            continue
        seen.add(key)
        phrase_scores.append((sum(word_score[w] for w in phrase), key))

    phrase_scores.sort(key=lambda x: x[0], reverse=True)
    return [phrase for _, phrase in phrase_scores[:max_keywords]]


def _detect_language(text: str) -> str:
    sample = text.strip()
    if len(sample) < 20:
        return "unknown"
    try:
        return detect(sample)
    except LangDetectException:
        return "unknown"


def _extract_fields(text: str, filename: str, file_type: str) -> dict:
    stripped = re.sub(r"\[.*?\]", "", text).strip()
    normalized = _normalize_text(stripped)
    word_count = len(normalized.split())

    summary = _summarize(normalized) if normalized else ""
    keywords = _rake_keywords(normalized)
    language = _detect_language(stripped)

    category_map = {
        "pdf": "document",
        "txt": "text",
        "csv": "data",
        "json": "data",
        "md": "documentation",
        "docx": "document",
    }
    category = category_map.get(file_type, "other")

    base = os.path.splitext(filename)[0]
    title = base.replace("_", " ").replace("-", " ").title()

    return {
        "title": title,
        "category": category,
        "summary": summary,
        "keywords": keywords,
        "word_count": word_count,
        "language": language,
    }


class BaseTask(Task):
    abstract = True

    def on_failure(self, exc, task_id, args, kwargs, einfo):
        job_id = args[0] if args else None
        if not job_id:
            return
        try:
            with SyncSession() as session:
                _update_job(
                    session, job_id,
                    status=JobStatus.FAILED,
                    error_message=str(exc)[:500],
                    current_stage="failed",
                )
        except Exception as e:
            logger.error("on_failure DB update failed for job %s: %s", job_id, e)
        try:
            _emit(job_id, "job_failed", 0, "failed", str(exc)[:200])
        except Exception as e:
            logger.error("on_failure emit failed for job %s: %s", job_id, e)


@celery_app.task(
    bind=True,
    base=BaseTask,
    name="app.workers.tasks.process_document_task",
    max_retries=3,
    default_retry_delay=60,
)
def process_document_task(
    self,
    job_id: str,
    document_id: str,
    storage_path: str,
    original_filename: str,
    file_type: str,
) -> dict:
    session: Session | None = None
    try:
        session = SyncSession()

        _update_job(session, job_id, status=JobStatus.PROCESSING, progress=5, current_stage="started")
        _emit(job_id, "job_started", 5, "started", "Processing started")
        time.sleep(0.5)

        _update_job(session, job_id, progress=20, current_stage="parsing")
        _emit(job_id, "document_parsing_started", 20, "parsing", "Parsing document...")
        time.sleep(1.0)

        extracted_text = _extract_text_from_file(storage_path, file_type, session, document_id)

        _update_job(session, job_id, progress=45, current_stage="parsing_done")
        _emit(job_id, "document_parsing_completed", 45, "parsing_done", "Parsing complete")
        time.sleep(0.5)

        _update_job(session, job_id, progress=60, current_stage="extracting")
        _emit(job_id, "field_extraction_started", 60, "extracting", "Extracting fields...")
        time.sleep(1.0)

        fields = _extract_fields(extracted_text, original_filename, file_type)

        _update_job(session, job_id, progress=80, current_stage="extraction_done")
        _emit(job_id, "field_extraction_completed", 80, "extraction_done", "Fields extracted")
        time.sleep(0.5)

        _update_job(session, job_id, progress=90, current_stage="storing")
        _emit(job_id, "storing_result", 90, "storing", "Storing result...")

        result = session.query(ProcessingResult).filter(
            ProcessingResult.job_id == job_id
        ).first()

        if not result:
            result = ProcessingResult(job_id=uuid.UUID(job_id))
            session.add(result)

        result.title = fields["title"]
        result.category = fields["category"]
        result.summary = fields["summary"]
        result.keywords = fields["keywords"]
        result.word_count = fields["word_count"]
        result.language = fields["language"]
        result.extracted_text = extracted_text[:5000]
        result.raw_json = {
            "fields": fields,
            "metadata": {
                "filename": original_filename,
                "file_type": file_type,
                "storage_path": storage_path,
                "processed_at": datetime.now(timezone.utc).isoformat(),
            },
        }

        document = session.query(Document).filter(Document.id == document_id).first()
        if document and document.file_content is not None:
            document.file_content = None

        session.commit()
        time.sleep(0.5)

        _update_job(
            session, job_id,
            status=JobStatus.COMPLETED,
            progress=100,
            current_stage="completed",
            completed_at=datetime.now(timezone.utc),
        )
        _emit(job_id, "job_completed", 100, "completed", "Processing complete!")

        return {"job_id": job_id, "status": "completed", "fields": fields}

    except Exception as exc:
        if session:
            try:
                session.rollback()
            except Exception:
                pass

        if isinstance(exc, MaxRetriesExceededError):
            raise

        logger.warning(
            "Task failed for job %s (attempt %d/%d): %s",
            job_id, self.request.retries + 1, self.max_retries + 1, exc,
        )
        raise self.retry(exc=exc, countdown=30)

    finally:
        if session:
            try:
                session.close()
            except Exception:
                pass